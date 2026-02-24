# Main GUI window for Mat Model Lab

import os
import sys
import datetime
try:
    import winreg
except ImportError:
    winreg = None

from mml_utils.paths import resource_path

import numpy as np
from PyQt6.QtWidgets import (QMainWindow, QApplication, QWidget, QPushButton, 
                             QVBoxLayout, QHBoxLayout, QLabel, QComboBox, 
                             QFileDialog, QTabWidget, QGroupBox, QRadioButton,
                             QCheckBox, QSpinBox, QDoubleSpinBox, QMessageBox,
                              QLineEdit, QGridLayout, QFrame)
from PyQt6.QtGui import QIcon, QPixmap, QAction, QPalette, QColor, QPainter, QPen, QTextDocument
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtCore import Qt



from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import matplotlib
import matplotlib.cm as cm
import matplotlib.colors as colors

# Absolute imports for stability in script execution
from core import Young_3D, Bulk_3D, Poisson_3D, Shear_3D, Hardness_3D, ElasticVRH3D
from core.shear import Shear_4D
from core.poisson import Poisson_4D
from core import check_stability_detailed
from visualization import ElasticPlot_2D, ElasticPlot_3D
from visualization.plot_slice import Plot_Slice_2D
from mml_utils.data_io import Elastic_Read
from mml_utils.report_generator import generate_report
from mml_utils.code_export import export_model
from core.crystal_type import list_crystal_types, CRYSTAL_TYPES_3D, CRYSTAL_TYPES_2D, get_enabled_indices, fill_symmetric_matrix, identify_crystal_type

from gui.widgets.plot_canvas import MatplotlibCanvas
from gui.widgets.help_dialog import HelpDialog
import PyQt6.QtCore as QtCore
from PyQt6.QtWidgets import (QMainWindow, QApplication, QWidget, QPushButton, 
                             QVBoxLayout, QHBoxLayout, QLabel, QComboBox, 
                             QFileDialog, QTabWidget, QGroupBox, QRadioButton,
                             QDoubleSpinBox, QCheckBox, QStackedWidget, QGridLayout,
                             QScrollArea, QMessageBox, QSpinBox) # Ensure QEvent is not strictly needed if we override visible methods, but showEvent signature uses event


class ElasticityWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._is_updating_cij = False
        self.init_ui()

    def showEvent(self, event):
        """Handle widget show event to add status bar widget"""
        super().showEvent(event)
        if self.window().statusBar():
            # Check if likely already added to avoid duplicates if re-shown?
            # removeWidget first to be safe or check parent?
            # statusBar takes ownership, so removing it first is safe.
            self.window().statusBar().removeWidget(self.stability_label) 
            self.window().statusBar().addPermanentWidget(self.stability_label)
            self.stability_label.show()
            
    def hideEvent(self, event):
        """Handle widget hide event to remove status bar widget"""
        super().hideEvent(event)
        if self.window().statusBar():
            self.window().statusBar().removeWidget(self.stability_label)
    
    def loadMaterialFromDatabase(self, material):
        """Load material data from the database browser.
        
        Parameters
        ----------
        material : dict
            Material data dictionary with keys: name, crystal_type, cij, source, description
        """
        import numpy as np
        
        # Get material properties
        name = material.get('name', 'Unknown')
        crystal_type = material.get('crystal_type', 'Cubic')
        cij_data = material.get('cij', [])
        dimension = material.get('dimension', '3D')
        
        # Convert to numpy array
        cij_array = np.array(cij_data)
        
        if cij_array.shape != (6, 6):
            print(f"Warning: Invalid Cij shape {cij_array.shape}, expected (6, 6)")
            return
        
        # Set dimension
        if hasattr(self, 'dimension_combo'):
            dim_text = '2D Materials' if dimension == '2D' else '3D Materials'
            idx = self.dimension_combo.findText(dim_text)
            if idx >= 0:
                self.dimension_combo.setCurrentIndex(idx)
        
        # Set crystal type
        idx = self.crystal_type_combo.findText(crystal_type)
        if idx >= 0:
            self.crystal_type_combo.blockSignals(True)
            self.crystal_type_combo.setCurrentIndex(idx)
            self.crystal_type_combo.blockSignals(False)
            self.onCrystalTypeChanged(idx)
        
        # Populate Cij matrix
        self._is_updating_cij = True
        for i in range(6):
            for j in range(6):
                if (i, j) in self.cij_inputs:
                    self.cij_inputs[(i, j)].setValue(cij_array[i, j])
        self._is_updating_cij = False
        
        # Update internal state
        self.cij = cij_array
        self.data_name = name
        
        # Set data mode to "From Database"
        idx = self.data_mode_combo.findText('From Database')
        if idx >= 0:
            self.data_mode_combo.setCurrentIndex(idx)
        
        print(f"Loaded material: {name} ({crystal_type})")
    def onDataModeChanged(self, index):
        """Handle changes in the data input mode
        
        Parameters
        ----------
        index : int
            Index of the selected data mode in the combo box
            0: Manual Input
            1: From File
            2: From Database
        """
        # Get the selected data mode
        data_mode = self.data_mode_combo.currentText()
        print(f"Data mode changed to: {data_mode}")
        
        # Update UI based on the selected data mode
        if data_mode == 'Manual Input':
            # Show manual input controls
            self.crystal_type_combo.setEnabled(True)
            self.crystal_type_label.setVisible(True)
            self.crystal_type_combo.setVisible(True)
            self.elastic_constants_group.setVisible(True)
            
            # Apply current crystal type settings
            # Reset matrix display when switching to Manual Input for a fresh start
            for (i, j), input_field in self.cij_inputs.items():
                input_field.setValue(0.0)
                
            self.onCrystalTypeChanged(self.crystal_type_combo.currentIndex())
            
        elif data_mode == 'From File':
            # Keep matrix visible - same layout as Manual Input
            self.crystal_type_combo.setEnabled(False)  # Disable since we auto-detect
            self.crystal_type_label.setVisible(True)
            self.crystal_type_combo.setVisible(True)
            self.elastic_constants_group.setVisible(True)
            
            # Immediately open file dialog
            self.loadElasticConstantsFromFile()
            
        elif data_mode == 'From Database':
            # Keep layout the same as Manual Input - no layout changes
            self.crystal_type_combo.setEnabled(False)  # Disable since we'll set from database
            self.crystal_type_label.setVisible(True)
            self.crystal_type_combo.setVisible(True)
            self.elastic_constants_group.setVisible(True)
            
            # Open the Material Browser dialog
            from gui.widgets.material_browser import MaterialBrowser
            dialog = MaterialBrowser(self, filter_category='elastic')
            dialog.materialSelected.connect(self.loadMaterialFromDatabase)
            result = dialog.exec()
            
            # If user cancelled, revert to Manual Input
            if result == 0:  # QDialog.Rejected
                self.data_mode_combo.blockSignals(True)
                self.data_mode_combo.setCurrentIndex(0)  # Manual Input
                self.data_mode_combo.blockSignals(False)
                self.onDataModeChanged(0)  # Apply Manual Input settings
            
    def onCrystalTypeChanged(self, index):
        """Handle changes in the crystal type selection
        
        Parameters
        ----------
        index : int
            Index of the selected crystal type in the combo box
            0: Cubic
            1: Hexagonal
            2: Tetragonal
            3: Orthorhombic
            4: Monoclinic
            5: Triclinic
        """
        # Get the selected crystal type
        crystal_type = self.crystal_type_combo.currentText()
        self.crystal_type = crystal_type
        
        print(f"Crystal type changed to: {crystal_type}")
        
        # Reset input values to 0.0 when crystal type changes manually
        # Note: If called programmatically during file load, rely on blocking signals
        self._is_updating_cij = True # Block recursive updates during reset
        for (i, j), widget in self.cij_inputs.items():
             widget.setValue(0.0)
        self._is_updating_cij = False
        
        # If placeholder is selected, disable all inputs
        if crystal_type == '-- Select Crystal Type --':
            # In 2D mode, apply dimension constraints first
            is_3d = True
            if hasattr(self, 'dimension_combo'):
                is_3d = self.dimension_combo.currentText() == '3D Materials'
            self._update_matrix_for_dimension(is_3d)
            return
        
        # Determine if 3D or 2D mode
        is_3d = True
        if hasattr(self, 'dimension_combo'):
            is_3d = self.dimension_combo.currentText() == '3D Materials'
        
        # Find crystal type index by name
        crystal_type_id = 0
        crystal_types_dict = CRYSTAL_TYPES_3D if is_3d else CRYSTAL_TYPES_2D
        for c_id, c_info in crystal_types_dict.items():
            if c_info['name'] == crystal_type:
                crystal_type_id = c_id
                break
                
        # Default to Triclinic/Oblique logic if unknown
        if crystal_type_id == 0 and crystal_type != '-- Select Crystal Type --':
            crystal_type_id = 10 if is_3d else 4  # Triclinic for 3D, Oblique for 2D
             
        # Determine Independent, Dependent, and Zero cells using a probe
        # 1. Get Independents
        visible_constants = get_enabled_indices(crystal_type_id, is_3d=is_3d)
        independent_set = set(visible_constants)
        
        # 2. Probe for Dependents (Non-zero but not independent)
        non_zero_dependents = set()
        if crystal_type_id > 0:
            C_probe = np.zeros((6, 6))
            # Set independents to unique non-zero values to ensure dependencies trigger
            # Avoid 0 or 1 to minimize accidental matches, use 10+idx
            for idx, (i, j) in enumerate(visible_constants):
                val = 10.0 + idx
                C_probe[i, j] = val
                C_probe[j, i] = val
            
            C_filled = fill_symmetric_matrix(C_probe, crystal_type_id, is_3d=is_3d)
            
            for i in range(6):
                for j in range(i, 6): # Upper triangle
                    if (i, j) not in independent_set and abs(C_filled[i, j]) > 1e-6:
                        non_zero_dependents.add((i, j))

        # For 2D mode, define which cells are allowed
        enabled_2d_cells = {(0, 0), (0, 1), (1, 1), (0, 5), (1, 5), (5, 5)}
        
        # First, disable all inputs and apply dimension-based styling
        for (i, j), input_field in self.cij_inputs.items():
            input_field.setEnabled(False)
            # In 2D mode, cells outside 2D range should be black/disabled
            if not is_3d and (i, j) not in enabled_2d_cells:
                input_field.setStyleSheet("background-color: #444; color: #666;")
            else:
                input_field.setStyleSheet("")
        
        # Enable Independent inputs (Standard Style) - only if allowed by dimension
        for i, j in visible_constants:
            if (i, j) in self.cij_inputs:
                # In 2D mode, only enable if it's a valid 2D cell
                if is_3d or (i, j) in enabled_2d_cells:
                    self.cij_inputs[(i, j)].setEnabled(True)
                    self.cij_inputs[(i, j)].setStyleSheet("")
                
        # Style Dependent inputs (Read-Only but Active-looking)
        for i, j in non_zero_dependents:
            if (i, j) in self.cij_inputs:
                # In 2D mode, skip styling cells outside 2D range
                if not is_3d and (i, j) not in enabled_2d_cells:
                    continue
                # Keep Enabled=False (Read-Only) but change style to look "Linked"
                # Using a light blue/green tint to show it's non-zero/calculated
                self.cij_inputs[(i, j)].setStyleSheet("""
                    QDoubleSpinBox {
                        background-color: #e0f7fa; 
                        color: #006064;
                        border: 1px dashed #4dd0e1;
                    }
                """)

    def onCijChanged(self, value):
        """Auto-sync dependent elastic constants based on symmetry rules."""
        # Only strict sync in Manual Input mode to prevent overwriting file data
        if self._is_updating_cij or self.data_mode_combo.currentText() != 'Manual Input':
            return

        self._is_updating_cij = True
        try:
            # 1. Read current matrix from GUI
            C = np.zeros((6, 6))
            for (i, j), widget in self.cij_inputs.items():
                val = widget.value()
                C[i, j] = val
                C[j, i] = val
            
            # 2. Identify Crystal Type ID
            crystal_type = self.crystal_type_combo.currentText()
            
            is_3d = True
            if hasattr(self, 'dimension_combo'):
                is_3d = self.dimension_combo.currentText() == '3D Materials'
                
            crystal_type_id = 0
            types_dict = CRYSTAL_TYPES_3D if is_3d else CRYSTAL_TYPES_2D
            
            for c_id, c_info in types_dict.items():
                if c_info['name'] == crystal_type:
                    crystal_type_id = c_id
                    break
            
            if crystal_type_id == 0:
                print(f"Warning: Unknown crystal type '{crystal_type}'")
                return

            # 3. Apply Symmetry Rules
            C_filled = fill_symmetric_matrix(C, crystal_type_id, is_3d=is_3d)

            # 4. Update GUI (Dependent Values)
            for (i, j), widget in self.cij_inputs.items():
                # Update if value changed (avoids unnecessary painting)
                # We specifically want to update disabled/dependent fields
                # But updating enabled ones is also fine as long as logic is correct (Cij=Cji)
                new_val = C_filled[i, j] if j >= i else C_filled[i, j] # Matrix is symmetric
                if abs(widget.value() - new_val) > 1e-6:
                   widget.setValue(new_val)
                   
        except Exception as e:
            print(f"Error in auto-sync: {e}")
        finally:
            self._is_updating_cij = False

    def onPropertyChanged(self, index):
        """Handle changes in the display property dropdown - auto-replot if data exists."""
        # Only auto-replot if we have already calculated (cij is not all zeros)
        if hasattr(self, 'cij') and np.any(self.cij != 0):
            self.calculateProperties()
    
    def onVisModeChanged(self, checked):
        """Toggle Plane Selection enabled state based on 2D/3D mode."""
        # Enable Plane Selection only when 2D is selected (keeps layout stable)
        if hasattr(self, 'plane_group'):
            self.plane_group.setEnabled(checked)
        
        # If in 2D mode, check if we need to show arbitrary plane inputs
        if checked and hasattr(self, 'arbitrary_plane_radio'):
            is_arbitrary = self.arbitrary_plane_radio.isChecked()
            if hasattr(self, 'hkl_widgets'):
                for widget in self.hkl_widgets:
                    widget.setEnabled(is_arbitrary)

        # Toggle Toolbar Actions (Grid/Mesh/Transparency) - Disable in 2D, Enable in 3D
        # checked is True if 2D is selected
        is_3d = not checked
        if hasattr(self, 'grid_action'):
            self.grid_action.setEnabled(is_3d)
        if hasattr(self, 'mesh_action'):
            self.mesh_action.setEnabled(is_3d)
        if hasattr(self, 'transparency_action'):
            self.transparency_action.setEnabled(is_3d)
            
    def onPlaneTypeChanged(self):
        """Handle change between Standard and Arbitrary plane types"""
        if not hasattr(self, 'arbitrary_plane_radio'):
            return
            
        is_arbitrary = self.arbitrary_plane_radio.isChecked()
        
        # Toggle HKL inputs enabled state
        if hasattr(self, 'hkl_widgets'):
            for widget in self.hkl_widgets:
                widget.setEnabled(is_arbitrary)
        
        # Toggle standard plane checkboxes enabled state
        for cb in self.plane_checkboxes.values():
            cb.setEnabled(not is_arbitrary)
        
        # Auto-replot
        self.onPropertyChanged(0)

    def _populate_crystal_types(self):
        """Populate crystal type combo based on current dimension selection."""
        self.crystal_type_combo.clear()
        
        # Determine if 3D or 2D
        is_3d = True
        if hasattr(self, 'dimension_combo'):
            is_3d = self.dimension_combo.currentText() == '3D Materials'
        
        # Get types for current dimension
        types_dict = list_crystal_types(is_3d=is_3d)
        # Sort by ID to ensure consistent order
        sorted_types = [types_dict[i] for i in sorted(types_dict.keys())]
        self.crystal_type_combo.addItems(['-- Select Crystal Type --'] + sorted_types)
        
    def onDimensionChanged(self, text):
        """Handle dimension change between 3D and 2D materials."""
        print(f"Dimension changed to: {text}")
        is_3d = (text == '3D Materials')
        
        # Update crystal type options
        self._populate_crystal_types()
        
        # Update matrix input enabled states for 2D
        self._update_matrix_for_dimension(is_3d)
        
        # Update visualization options for 2D mode
        self._update_visualization_for_dimension(is_3d)
        
        # Clear current data
        self.cij = np.zeros((6, 6))
        
    def _update_matrix_for_dimension(self, is_3d):
        """Enable/disable matrix cells based on 2D/3D dimension."""
        if is_3d:
            # 3D: enable all upper triangle cells
            for (i, j), field in self.cij_inputs.items():
                field.setEnabled(True)
                field.setStyleSheet("")  # Clear any grey styling
        else:
            # 2D: only enable C11, C12, C22, C16, C26, C66
            # In 0-indexed: (0,0), (0,1), (1,1), (0,5), (1,5), (5,5)
            enabled_2d = {(0, 0), (0, 1), (1, 1), (0, 5), (1, 5), (5, 5)}
            for (i, j), field in self.cij_inputs.items():
                if (i, j) in enabled_2d:
                    field.setEnabled(True)
                    field.setStyleSheet("")
                else:
                    field.setEnabled(False)
                    field.setValue(0.0)
                    field.setStyleSheet("background-color: #444; color: #666;")
    
    def _update_visualization_for_dimension(self, is_3d):
        """Enable/disable visualization options based on 2D/3D dimension."""
        # Improved disabled style: Grey text only, let native/theme handle background transparency
        disabled_style = "QCheckBox { color: #888; } QCheckBox::indicator { background-color: #e0e0e0; border: 1px solid #ccc; }"
        enabled_style = ""
        
        # Disable 2D/3D plot toggle in 2D mode (only 2D polar plot makes sense)
        radio_disabled_style = "QRadioButton { color: #888; } QRadioButton::indicator { background-color: #ccc; border: 1px solid #aaa; }"
        if hasattr(self, 'vis_2d_radio'):
            self.vis_2d_radio.setEnabled(is_3d)
            self.vis_2d_radio.setStyleSheet("" if is_3d else radio_disabled_style)
            if not is_3d:
                self.vis_2d_radio.setChecked(True)  # Force 2D plot mode
        if hasattr(self, 'vis_3d_radio'):
            self.vis_3d_radio.setEnabled(is_3d)
            self.vis_3d_radio.setStyleSheet("" if is_3d else radio_disabled_style)
        
        # Disable plane selection in 2D mode (no z-axis)
        if hasattr(self, 'plane_group'):
            self.plane_group.setEnabled(is_3d)
            # Apply greyed out style to the whole group
            if not is_3d:
                self.plane_group.setStyleSheet("""
                    QGroupBox { color: #888; }
                    QGroupBox::title { color: #888; }
                    QRadioButton { color: #888; }
                    QRadioButton::indicator { background-color: #ccc; border: 1px solid #aaa; }
                    QCheckBox { color: #888; }
                    QCheckBox::indicator { background-color: #ccc; border: 1px solid #aaa; }
                    QLabel { color: #888; }
                    QDoubleSpinBox { background-color: #e0e0e0; color: #888; }
                """)
            else:
                self.plane_group.setStyleSheet("")
        
        # Disable 3D View Options in 2D mode
        if hasattr(self, 'view_group'):
            self.view_group.setEnabled(is_3d)
            if not is_3d:
                self.view_group.setStyleSheet("""
                    QGroupBox { color: #888; }
                    QGroupBox::title { color: #888; }
                    QPushButton { color: #888; }
                """)
            else:
                self.view_group.setStyleSheet("")
        
        # Update Display Property dropdown based on dimension
        if hasattr(self, 'display_property_combo'):
            current_text = self.display_property_combo.currentText()
            self.display_property_combo.blockSignals(True)
            self.display_property_combo.clear()
            if is_3d:
                # 3D properties
                self.display_property_combo.addItems([
                    "Young's Modulus (E)",
                    "Shear Modulus (G)",
                    "Poisson's Ratio (v)",
                    "Bulk Modulus (B)",
                    "Hardness (H)"
                ])
            else:
                # 2D properties (no Hardness for 2D materials)
                self.display_property_combo.addItems([
                    "Young's Modulus (E)",
                    "Shear Modulus (G)",
                    "Poisson's Ratio (v)",
                    "Bulk Modulus (B)"
                ])
            # Restore previous selection if valid
            idx = self.display_property_combo.findText(current_text)
            if idx >= 0:
                self.display_property_combo.setCurrentIndex(idx)
            self.display_property_combo.blockSignals(False)

        # Update Property Checkboxes (Properties to Calculate)
        if hasattr(self, 'property_checkboxes') and 'H' in self.property_checkboxes:
            hardness_cb = self.property_checkboxes['H']
            hardness_cb.setEnabled(is_3d)
            if not is_3d:
                hardness_cb.setChecked(False)
                hardness_cb.setStyleSheet(disabled_style)
            else:
                hardness_cb.setChecked(True) # Optional: re-check if enabling? Or just enable
                hardness_cb.setStyleSheet("")

    def showDocumentation(self):
        """Open the help documentation dialog"""
        # Help Dialog usually needs a window parent, self is a widget but inside a window
        dialog = HelpDialog(self)
        dialog.exec()
        

    
    def initializeData(self):
        """Initialize data structures for the application"""
        # Default elastic constants (6x6 matrix)
        self.cij = np.zeros((6, 6))
        
        # Default crystal type (None until selected)
        self.crystal_type = None
        
        # Default number of points for 3D plotting
        self.n_points = 100
        
        # Default properties to plot
        self.plot_properties = ['E', 'G', 'v', 'B', 'H']
        
        # Default property display mode (Average, Min, Max)
        self.property_mode = 'Ave'
        
        # Current data source
        self.data_source = None
        self.data_name = 'Untitled'
        
    def init_ui(self):
        # Initialize data structures
        self.initializeData()
        
        # Initialize UI widgets
        self.init_ui_components()
        
        # Apply theme specific settings for plot (local only)
        # self.update_theme('dark') # Parent should trigger this or default


    def set_3d_view(self, elev, azim):
        """Set the view angle for the 3D plot"""
        if hasattr(self, 'canvas') and self.canvas.axes:
            # Check if current plot is 3D
            if hasattr(self.canvas.axes, 'name') and self.canvas.axes.name == '3d':
                self.canvas.axes.view_init(elev=elev, azim=azim)
                self.canvas.draw()
            else:
                 # If likely 2D, warn or ignore?
                 # If user clicks while in 2D mode, we might want to switch to 3D? 
                 # But sticking to just changing view if 3D is safer.
                 pass

    def update_theme(self, theme='dark'):
        """Update widget visualization theme"""
        self.current_theme = theme
        
        # Plot Colors
        if theme == 'dark':
            text_color = "#f0f0f0" 
            plot_bg = "#ffffff" # Keep plot white as requested elsewhere
        else:
            text_color = "#333333"
            plot_bg = "#ffffff"
            
        if hasattr(self, 'canvas'):
            self.canvas.fig.set_facecolor(plot_bg)
            self.canvas.draw()
            
            # Re-create toolbar to update icon colors based on theme if needed
            # Remove old toolbar
            if hasattr(self, 'toolbar'):
                self.toolbar.setParent(None)
                self.toolbar.deleteLater()
            
            # Create new toolbar
            from matplotlib.backends.backend_qtagg import NavigationToolbar2QT as NavigationToolbar
            self.toolbar = NavigationToolbar(self.canvas, self)
            
            # We need to find where the toolbar was in the layout and insert it there.
            # self.canvas.parent() should be vis_area.
            vis_area = self.canvas.parent()
            if vis_area and vis_area.layout():
                vis_area.layout().insertWidget(0, self.toolbar)
                
            # Find the 'Save' action to insert after (usually the last icon)
            insert_before_action = None
            save_action = None
            for action in self.toolbar.actions():
                if action.toolTip() and 'Save' in action.toolTip():
                    save_action = action
                    break
            
            if save_action:
                actions = self.toolbar.actions()
                try:
                    idx = actions.index(save_action)
                    if idx + 1 < len(actions):
                        insert_before_action = actions[idx + 1]
                except ValueError:
                    pass

            # Create Actions
            icon_color = text_color
            
            self.grid_action = QAction(self)
            self.grid_action.setText("Grid") # Text for tooltip/overflow
            self.grid_action.setIcon(self._create_toolbar_icon('grid', icon_color))
            self.grid_action.setCheckable(True)
            self.grid_action.setChecked(getattr(self, 'show_grid', True))
            self.grid_action.setToolTip("Toggle Grid and Axis visibility")
            self.grid_action.triggered.connect(self.toggleGrid)

            self.mesh_action = QAction(self)
            self.mesh_action.setText("Mesh")
            self.mesh_action.setIcon(self._create_toolbar_icon('mesh', icon_color))
            self.mesh_action.setCheckable(True)
            self.mesh_action.setChecked(getattr(self, 'show_mesh', False))
            self.mesh_action.setToolTip("Toggle Surface Mesh lines")
            self.mesh_action.triggered.connect(self.toggleMesh)

            self.transparency_action = QAction(self)
            self.transparency_action.setText("Alpha")
            self.transparency_action.setIcon(self._create_toolbar_icon('alpha', icon_color))
            self.transparency_action.setCheckable(True)
            self.transparency_action.setChecked(getattr(self, 'show_transparent', False))  # Default opaque
            self.transparency_action.setToolTip("Toggle Surface Transparency")
            self.transparency_action.triggered.connect(self.toggleTransparency)

            # Insert or Add
            if insert_before_action:
                self.toolbar.insertSeparator(insert_before_action)
                self.toolbar.insertAction(insert_before_action, self.grid_action)
                self.toolbar.insertAction(insert_before_action, self.mesh_action)
                self.toolbar.insertAction(insert_before_action, self.transparency_action)
            else:
                self.toolbar.addSeparator()
                self.toolbar.addAction(self.grid_action)
                self.toolbar.addAction(self.mesh_action)
                self.toolbar.addAction(self.transparency_action)


    def init_ui_components(self):
        # Create central widget and main layout
        # self is the widget
        main_layout = QHBoxLayout(self)
        
        # Create control panel
        control_panel = QWidget()
        control_layout = QVBoxLayout(control_panel)
        control_panel.setMaximumWidth(490)
        
        # Data input section
        data_group = QGroupBox('Data Input')
        data_layout = QVBoxLayout(data_group)
        
        # Dimension selection (3D vs 2D materials)
        dimension_layout = QHBoxLayout()
        dimension_label = QLabel('Dimension:')
        self.dimension_combo = QComboBox()
        self.dimension_combo.addItems(['3D Materials', '2D Materials'])
        self.dimension_combo.currentTextChanged.connect(self.onDimensionChanged)
        dimension_layout.addWidget(dimension_label)
        dimension_layout.addWidget(self.dimension_combo)
        data_layout.addLayout(dimension_layout)
        
        # Data mode selection
        data_mode_layout = QHBoxLayout()
        data_mode_label = QLabel('Data Mode:')
        self.data_mode_combo = QComboBox()
        self.data_mode_combo.addItems(['Manual Input', 'From File', 'From Database'])
        # Use activated instead of currentIndexChanged so it triggers every time, even when same item is selected
        self.data_mode_combo.activated.connect(self.onDataModeChanged)
        data_mode_layout.addWidget(data_mode_label)
        data_mode_layout.addWidget(self.data_mode_combo)
        data_layout.addLayout(data_mode_layout)
        
        # Crystal type selection
        crystal_type_layout = QHBoxLayout()
        self.crystal_type_label = QLabel('Crystal Type:')
        self.crystal_type_combo = QComboBox()
        # Populate crystal types based on current dimension
        self._populate_crystal_types()
        # Use activated to force update even if selecting same
        self.crystal_type_combo.activated.connect(self.onCrystalTypeChanged)
        crystal_type_layout.addWidget(self.crystal_type_label)
        crystal_type_layout.addWidget(self.crystal_type_combo)
        data_layout.addLayout(crystal_type_layout)
        
        # Elastic constants input section with scroll area
        self.elastic_constants_group = QGroupBox('Elastic Stiffness Constants (Cij) (GPa)')
        elastic_group_layout = QVBoxLayout(self.elastic_constants_group)
        
        # Create a 6x6 matrix layout (proper matrix display)
        matrix_widget = QWidget()
        self.elastic_constants_layout = QGridLayout(matrix_widget)
        self.elastic_constants_layout.setVerticalSpacing(6)
        self.elastic_constants_layout.setHorizontalSpacing(8)
        
        # Add column headers (1-6)
        for j in range(6):
            header = QLabel(f'{j+1}')
            header.setAlignment(Qt.AlignmentFlag.AlignCenter)
            header.setStyleSheet("font-weight: bold; color: #888;")
            self.elastic_constants_layout.addWidget(header, 0, j+1)
        
        # Create input fields for elastic constants
        # Store both labels and inputs for easy visibility control
        self.cij_inputs = {}
        self.cij_labels = {}
        for i in range(6):
            # Row label
            row_label = QLabel(f'C{i+1}x')
            row_label.setStyleSheet("font-weight: bold; color: #888;")
            self.elastic_constants_layout.addWidget(row_label, i+1, 0)
            
            for j in range(6):
                if j >= i:  # Only create inputs for upper triangular part (due to symmetry)
                    input_field = QDoubleSpinBox()
                    input_field.setRange(-9999, 9999)
                    input_field.setDecimals(1)
                    input_field.setSingleStep(1.0)
                    input_field.setValue(0.0)
                    input_field.setMinimumWidth(55)
                    input_field.setMaximumWidth(70)
                    input_field.setButtonSymbols(QDoubleSpinBox.ButtonSymbols.NoButtons)
                    # Connect value change to auto-sync slot
                    input_field.valueChanged.connect(self.onCijChanged)
                    self.cij_inputs[(i, j)] = input_field
                    # Create hidden label for compatibility
                    label = QLabel(f'C{i+1}{j+1}:')
                    label.setVisible(False)
                    self.cij_labels[(i, j)] = label
                    self.elastic_constants_layout.addWidget(input_field, i+1, j+1)
                else:
                    # Lower triangle: mirror of upper
                    mirror_label = QLabel('—')
                    mirror_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    mirror_label.setStyleSheet("color: #555;")
                    self.elastic_constants_layout.addWidget(mirror_label, i+1, j+1)
        
        elastic_group_layout.addWidget(matrix_widget)
        
        # Add file selection button for 'From File' mode
        self.file_select_layout = QHBoxLayout()
        self.file_path_label = QLabel('No file selected')
        self.file_select_button = QPushButton('Select File')
        self.file_select_button.clicked.connect(self.selectFile)
        self.file_select_layout.addWidget(self.file_path_label)
        self.file_select_layout.addWidget(self.file_select_button)
        data_layout.addLayout(self.file_select_layout)
        
        # Initially hide file selection
        self.file_path_label.setVisible(False)
        self.file_select_button.setVisible(False)
        
        # Add elastic constants group to data layout
        data_layout.addWidget(self.elastic_constants_group)
        
        # Calculation options
        calc_group = QGroupBox('Calculation Options')
        calc_layout = QVBoxLayout(calc_group)
        
        # Number of points for 3D plotting
        points_layout = QHBoxLayout()
        points_label = QLabel('Number of Points:')
        self.points_spinbox = QSpinBox()
        self.points_spinbox.setRange(10, 1000)
        self.points_spinbox.setValue(self.n_points)
        self.points_spinbox.setSingleStep(10)
        points_layout.addWidget(points_label)
        points_layout.addWidget(self.points_spinbox)
        calc_layout.addLayout(points_layout)
        
        # Properties to calculate
        properties_group = QGroupBox('Properties to Calculate')
        properties_layout = QGridLayout(properties_group)
        
        self.property_checkboxes = {}
        idx = 0
        for prop, label in zip(['E', 'G', 'v', 'B', 'H'], 
                             ['Young\'s Modulus', 'Shear Modulus', 'Poisson\'s Ratio', 'Bulk Modulus', 'Hardness']):
            checkbox = QCheckBox(label)
            checkbox.setChecked(prop in self.plot_properties)
            self.property_checkboxes[prop] = checkbox
            row, col = divmod(idx, 2)
            properties_layout.addWidget(checkbox, row, col)
            idx += 1
        
        calc_layout.addWidget(properties_group)
        
        # Calculate button
        self.calculate_button = QPushButton('Calculate')
        self.calculate_button.setObjectName('calc_btn')
        self.calculate_button.clicked.connect(self.calculateProperties)
        calc_layout.addWidget(self.calculate_button)
        
        self.calculate_button.setObjectName('calc_btn')
        self.calculate_button.clicked.connect(self.calculateProperties)
        calc_layout.addWidget(self.calculate_button)
        
        # Export Button REMOVED from here
        
        # Add calculation group to control layout
        control_layout.addWidget(data_group)
        control_layout.addWidget(calc_group)
        
        # Visualization options
        vis_group = QGroupBox('Visualization Options')
        vis_layout = QVBoxLayout(vis_group)
        

        
        # Property to display selector
        prop_select_layout = QHBoxLayout()
        prop_select_label = QLabel('Display Property:')
        self.display_property_combo = QComboBox()
        self.display_property_combo.addItems([
            "Young's Modulus (E)",
            "Shear Modulus (G)",
            "Poisson's Ratio (v)",
            "Bulk Modulus (B)",
            "Hardness (H)"
        ])
        # Store mapping of display names to property codes
        self.property_code_map = {
            "Young's Modulus (E)": 'E',
            "Shear Modulus (G)": 'G',
            "Poisson's Ratio (v)": 'v',
            "Bulk Modulus (B)": 'B',
            "Hardness (H)": 'H'
        }
        prop_select_layout.addWidget(prop_select_label)
        prop_select_layout.addWidget(self.display_property_combo)
        # Auto-update plot when property changes
        self.display_property_combo.currentIndexChanged.connect(self.onPropertyChanged)
        vis_layout.addLayout(prop_select_layout)

        # 2D/3D selection (Moved below Property Selector)
        vis_type_layout = QHBoxLayout()
        self.vis_2d_radio = QRadioButton('2D Plot')
        self.vis_3d_radio = QRadioButton('3D Plot')
        self.vis_3d_radio.setChecked(True)
        # Auto-update plot when switching 2D/3D mode
        self.vis_2d_radio.toggled.connect(self.onPropertyChanged)
        self.vis_3d_radio.toggled.connect(self.onPropertyChanged)
        # Also toggle Plane Selection visibility
        self.vis_2d_radio.toggled.connect(self.onVisModeChanged)
        vis_type_layout.addWidget(self.vis_2d_radio)
        vis_type_layout.addWidget(self.vis_3d_radio)
        vis_layout.addLayout(vis_type_layout)
        
        # 2D plane selection (for 2D plots only)
        self.plane_group = QGroupBox('Plane Selection')
        plane_grouplayout = QGridLayout(self.plane_group)
        # Add some spacing
        plane_grouplayout.setHorizontalSpacing(40)
        plane_grouplayout.setVerticalSpacing(8)
        
        # --- Row 0: Radio Buttons ---
        self.standard_plane_radio = QRadioButton("Standard")
        self.standard_plane_radio.setChecked(True)
        self.standard_plane_radio.toggled.connect(self.onPlaneTypeChanged)
        plane_grouplayout.addWidget(self.standard_plane_radio, 0, 0)
        
        self.arbitrary_plane_radio = QRadioButton("Arbitrary (hkl)")
        self.arbitrary_plane_radio.toggled.connect(self.onPlaneTypeChanged)
        plane_grouplayout.addWidget(self.arbitrary_plane_radio, 0, 1) # Column 1

        
        # --- Rows 1-3: Checkboxes (Left) and Inputs (Right) ---
        self.plane_checkboxes = {}
        self.hkl_inputs = []
        self.hkl_widgets = [] # For enabling/disabling
        
        # Standard Planes (Column 0)
        planes = ['xy', 'xz', 'yz']
        for i, plane in enumerate(planes):
            # Checkbox
            checkbox = QCheckBox(f'{plane.upper()} Plane')
            checkbox.setChecked(True)
            checkbox.stateChanged.connect(self.onPropertyChanged)
            self.plane_checkboxes[plane] = checkbox
            plane_grouplayout.addWidget(checkbox, i+1, 0)
            
            # HKL Input (Column 1)
            # Use a nested layout to keep label and input close together
            hkl_row_widget = QWidget()
            hkl_row_widget.setStyleSheet("background-color: transparent;")
            hkl_row_layout = QHBoxLayout(hkl_row_widget)
            hkl_row_layout.setContentsMargins(0, 0, 0, 0)
            hkl_row_layout.setSpacing(5)
            
            label_text = ['h', 'k', 'l'][i]
            lbl = QLabel(f'{label_text}:')
            lbl.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
            
            spin = QDoubleSpinBox()
            spin.setRange(-20, 20)
            spin.setDecimals(1)
            spin.setSingleStep(0.5)
            # Default to [1 1 1]
            spin.setValue(1.0)
            spin.valueChanged.connect(lambda _: self.onPropertyChanged(0))
            
            # Add to nested layout
            hkl_row_layout.addWidget(lbl)
            hkl_row_layout.addWidget(spin)
            hkl_row_layout.addStretch() # Push to left
            
            # Add container to main grid
            plane_grouplayout.addWidget(hkl_row_widget, i+1, 1)
            
            # Store references
            self.hkl_inputs.append(spin)
            self.hkl_widgets.append(lbl)
            self.hkl_widgets.append(spin)

            
        # Initial state for HKL widgets (Disabled)
        for widget in self.hkl_widgets:
            widget.setEnabled(False)


        # Initially disabled since 3D is default
        self.plane_group.setEnabled(False) 
        
        # --- 3D View Options (New Right Column) ---
        self.view_group = QGroupBox('3D View')
        view_layout = QGridLayout(self.view_group)
        view_layout.setHorizontalSpacing(10)
        view_layout.setVerticalSpacing(10)
        
        # View Buttons
        btn_default = QPushButton("Isometric")
        btn_default.clicked.connect(lambda: self.set_3d_view(30, -60)) # Matplotlib default-ish
        
        btn_x = QPushButton("X Direction")
        btn_x.clicked.connect(lambda: self.set_3d_view(0, 0)) # Elevation 0, Azimuth 0 -> View from X? Need to verify angles. 
        # Usually Azim=0, Elev=0 views along X axis? Or Y? 
        # In Matplotlib: 
        # azim=0, elev=0 -> View from +X? NO.
        # Actually: azim is rotation around Z. elev is angle from XY plane.
        # Default azim=-60, elev=30.
        # Let's test angles. "View X" usually means looking FROM X towards origin.
        # If I look from +X, Azim=0? 
        # Let's try standard presets:
        # X: elev=0, azim=0
        # Y: elev=0, azim=90
        # Z: elev=90, azim=0 (Top view)
        
        btn_y = QPushButton("Y Direction")
        btn_y.clicked.connect(lambda: self.set_3d_view(0, 90))
        
        btn_z = QPushButton("Z Direction")
        btn_z.clicked.connect(lambda: self.set_3d_view(90, -90)) # Top view, often azim=-90 puts Y up?
        
        # Add to layout
        view_layout.addWidget(btn_default, 0, 0)
        view_layout.addWidget(btn_z, 0, 1)
        view_layout.addWidget(btn_x, 1, 0)
        view_layout.addWidget(btn_y, 1, 1)
        
        
        # Split Layout: Left (Plane) - Right (View)
        split_layout = QHBoxLayout()
        split_layout.addWidget(self.plane_group, stretch=3) # Give Plane selection more space if needed, or equal?
        split_layout.addWidget(self.view_group, stretch=2)
        
        vis_layout.addLayout(split_layout)
        
        # Add visualization group to control layout
        control_layout.addWidget(vis_group)
        
        # Add control panel to main layout
        main_layout.addWidget(control_panel)
        
        # Create visualization area
        vis_area = QWidget()
        vis_layout = QVBoxLayout(vis_area)
        
        # Create matplotlib canvas for plotting
        self.canvas = MatplotlibCanvas(self, width=5, height=4, dpi=100)
        self.toolbar = NavigationToolbar(self.canvas, self)
        
        vis_layout.addWidget(self.toolbar)
        vis_layout.addWidget(self.canvas)
        
        # Add visualization area to main layout
        main_layout.addWidget(vis_area)
        


        # Status Message (Log to console or use a local label if needed)
        print("Elasticity Module Ready")

        # Stability indicator (Add to control panel instead of status bar)
        self.stability_label = QLabel("Structure Stability: Unknown")
        self.stability_label.setStyleSheet("font-weight: bold; color: #555; padding-right: 10px;") # Added padding for status bar
        self.stability_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Add to layout REMOVED - will add to status bar in showEvent
        
        # Add Export and Report buttons to bottom of control layout
        buttons_layout = QHBoxLayout()
        
        self.export_button = QPushButton('Export Model')
        self.export_button.clicked.connect(self.exportData)
        buttons_layout.addWidget(self.export_button)
        
        self.report_button = QPushButton('Generate Report')
        self.report_button.clicked.connect(lambda: self.generateReport('html'))
        buttons_layout.addWidget(self.report_button)
        
        control_layout.addSpacing(10)
        control_layout.addLayout(buttons_layout)
        # Initialize input states based on current crystal type selection (disabled by default)
        self.onCrystalTypeChanged(0)
    
    def selectFile(self):
        """Open file dialog to select elastic constants file (legacy method)"""
        self.loadElasticConstantsFromFile()
    
    def loadElasticConstantsFromFile(self):
        """Open file dialog, load elastic constants, populate matrix, and auto-detect crystal type"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, 'Open Elastic Constants File', '',
            'All Files (*);;Text Files (*.txt);;Excel Files (*.xlsx);;MATLAB Files (*.mat)')
        
        if not file_path:
            # User cancelled - don't change anything
            return
        
        try:
            # Read elastic constants from file
            # Elastic_Read returns (Cij, ComName, State) where Cij has shape (6, 6, n)
            cij_array, material_names, state = Elastic_Read(file_path)
            
            if state != 'OK':
                QMessageBox.warning(self, 'Warning', f'File loaded with warning: {state}')
            
            if cij_array is None or len(cij_array.shape) < 2:
                QMessageBox.critical(self, 'Error', 'Invalid data format. Expected stiffness matrix.')
                self.data_mode_combo.setCurrentIndex(0)
                return
            
            # Use the first material if multiple are present
            if len(cij_array.shape) == 3:
                cij_data = cij_array[:, :, 0]
                if cij_array.shape[2] > 1:
                    QMessageBox.information(self, 'Info', 
                        f'File contains {cij_array.shape[2]} materials. Using first: {material_names[0] if material_names else "Unknown"}')
            else:
                cij_data = cij_array
            
            # Store the loaded data
            self.cij = cij_data
            self.data_source = file_path
            
            # Populate the input matrix with loaded values
            for i in range(6):
                for j in range(6):
                    if j >= i and (i, j) in self.cij_inputs:
                        self.cij_inputs[(i, j)].setValue(self.cij[i, j])
            
            # Check if 3D or 2D mode
            is_3d = True
            if hasattr(self, 'dimension_combo'):
                is_3d = self.dimension_combo.currentText() == '3D Materials'
            
            # For 2D mode, define which cells are allowed
            enabled_2d_cells = {(0, 0), (0, 1), (1, 1), (0, 5), (1, 5), (5, 5)}
            
            # Auto-detect crystal type from the matrix pattern
            # Use the robust core detection logic
            if is_3d:
                type_id, detected_type = identify_crystal_type(self.cij)
            else:
                # For 2D, use a simpler detection based on which cells are non-zero
                # Check pattern to determine 2D crystal type
                detected_type = 'Oblique'  # Default to most general
                c11, c12, c22, c16, c26, c66 = self.cij[0,0], self.cij[0,1], self.cij[1,1], self.cij[0,5], self.cij[1,5], self.cij[5,5]
                
                # Hexagonal 2D: C11=C22, C66=(C11-C12)/2, C16=C26=0
                if abs(c16) < 1e-6 and abs(c26) < 1e-6:
                    if abs(c11 - c22) < 1e-6 and abs(c66 - (c11 - c12)/2) < 1e-6:
                        detected_type = 'Hexagonal'
                    elif abs(c11 - c22) < 1e-6:
                        detected_type = 'Square'
                    else:
                        detected_type = 'Rectangular'
                type_id = 0
                for c_id, c_info in CRYSTAL_TYPES_2D.items():
                    if c_info['name'] == detected_type:
                        type_id = c_id
                        break
            
            # Set crystal type in combo box by text
            index = self.crystal_type_combo.findText(detected_type)
            
            # Block signals to prevent onCrystalTypeChanged from clearing the data we just loaded
            self.crystal_type_combo.blockSignals(True)
            try:
                if index >= 0:
                    self.crystal_type_combo.setCurrentIndex(index)
                else:
                    # Default to Triclinic (3D) or Oblique (2D) if not found
                    fallback = 'Triclinic' if is_3d else 'Oblique'
                    fallback_index = self.crystal_type_combo.findText(fallback)
                    if fallback_index >= 0:
                         self.crystal_type_combo.setCurrentIndex(fallback_index)
                
                # Re-detect type_id after setting combo
                crystal_types_dict = CRYSTAL_TYPES_3D if is_3d else CRYSTAL_TYPES_2D
                type_id = 0
                for c_id, c_info in crystal_types_dict.items():
                    if c_info['name'] == detected_type:
                        type_id = c_id
                        break
                if type_id == 0: 
                    type_id = 10 if is_3d else 4  # Fallback
                
                visible_constants = get_enabled_indices(type_id, is_3d=is_3d)
                
                # Disable all first and apply dimension-based styling
                for (i, j), input_field in self.cij_inputs.items():
                    input_field.setEnabled(False)
                    if not is_3d and (i, j) not in enabled_2d_cells:
                        input_field.setStyleSheet("background-color: #444; color: #666;")
                    else:
                        input_field.setStyleSheet("")
                
                # Enable relevant ones and set read-only style
                for i, j in visible_constants:
                    if (i, j) in self.cij_inputs:
                        # In 2D mode, only enable if it's a valid 2D cell
                        if is_3d or (i, j) in enabled_2d_cells:
                            self.cij_inputs[(i, j)].setEnabled(True)
                            self.cij_inputs[(i, j)].setStyleSheet("""
                                QDoubleSpinBox {
                                    background-color: #e6f3ff;
                                    color: #333;
                                    border: 1px solid #4a90d9;
                                    border-radius: 3px;
                                    padding: 2px;
                                }
                            """)

            finally:
                self.crystal_type_combo.blockSignals(False)
            
            # self.crystal_type = detected_type # Sync internal state
            # self.crystal_type = detected_type # Sync internal state
            if self.window().statusBar():
                self.window().statusBar().showMessage(f'Loaded: {os.path.basename(file_path)} - Detected: {detected_type}')
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, 'Error', f'Failed to load file: {str(e)}')
            self.data_mode_combo.setCurrentIndex(0)
    
    # _detectCrystalType removed - replaced by core.crystal_type.identify_crystal_type
    
    def calculateProperties(self):
        """Calculate elastic properties based on input data"""
        # Show progress indicator and disable button to prevent multiple clicks
        self.calculate_button.setEnabled(False)
        self.calculate_button.setText('Calculating...')
        if self.window().statusBar():
            self.window().statusBar().showMessage('Calculating... Please wait.')
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        QApplication.processEvents()  # Force UI update
        
        try:
            self._doCalculation()
        finally:
            # Always restore UI state
            QApplication.restoreOverrideCursor()
            self.calculate_button.setEnabled(True)
            self.calculate_button.setText('Calculate')
            if self.window().statusBar():
                self.window().statusBar().showMessage('Ready')
    
    def _doCalculation(self):
        """Internal method that performs the actual calculation."""
        # Get the current data mode
        data_mode = self.data_mode_combo.currentText()
        
        if data_mode == 'Manual Input':
            # Get elastic constants from input fields
            for i in range(6):
                for j in range(6):
                    if j >= i:  # Upper triangular part
                        value = self.cij_inputs.get((i, j), QDoubleSpinBox()).value()
                        self.cij[i, j] = value
                        self.cij[j, i] = value  # Ensure symmetry
            
            # Fill dependent constants based on crystal symmetry
            self._fillCijBySymmetry()
            
        elif data_mode == 'From File':
            if not self.data_source:
                QMessageBox.warning(self, 'Warning', 'Please select a file first.')
                return
            # TODO: Load data from file using Elastic_Read
        elif data_mode == 'From Database':
            # Material data was already loaded via loadMaterialFromDatabase
            # self.cij is already populated, just proceed with calculation
            pass
        
        # Get selected properties to calculate
        properties = [prop for prop, checkbox in self.property_checkboxes.items() if checkbox.isChecked()]
        if not properties:
            QMessageBox.warning(self, 'Warning', 'Please select at least one property to calculate.')
            return
        
        # Get number of points for plotting
        n_points = self.points_spinbox.value()
        
        # Determine visualization type (2D or 3D)
        vis_type = '2D' if self.vis_2d_radio.isChecked() else '3D'
        
        # Get selected planes for 2D plotting
        planes = [plane for plane, checkbox in self.plane_checkboxes.items() if checkbox.isChecked()]
        if vis_type == '2D' and not planes:
            QMessageBox.warning(self, 'Warning', 'Please select at least one plane for 2D plotting.')
            return
        
        # Clear the current figure
        self.canvas.fig.clear()
        
        try:
            # Check if we're in 2D mode
            is_2d_material = False
            if hasattr(self, 'dimension_combo'):
                is_2d_material = self.dimension_combo.currentText() == '2D Materials'
            
            if is_2d_material:
                # For 2D materials, extract 3x3 submatrix (indices 0,1,5 for C11,C12,C22,C16,C26,C66)
                # The 2D stiffness matrix is:
                # | C11 C12 C16 |
                # | C12 C22 C26 |
                # | C16 C26 C66 |
                indices_2d = [0, 1, 5]
                C_2d = np.array([[self.cij[i, j] for j in indices_2d] for i in indices_2d])
                
                try:
                    S_2d = np.linalg.inv(C_2d)
                except np.linalg.LinAlgError:
                    QMessageBox.critical(self, 'Error', '2D Stiffness matrix is singular! Check your input values.')
                    return
                
                # Get selected property from Display Property combo
                selected_prop = self.display_property_combo.currentText() if hasattr(self, 'display_property_combo') else "Young's Modulus (E)"
                
                # For 2D materials, we'll do a simple polar plot
                theta_2d = np.linspace(0, 2*np.pi, 360)
                
                # Extract compliance components
                S11, S12, S16 = S_2d[0, 0], S_2d[0, 1], S_2d[0, 2]
                S22, S26, S66 = S_2d[1, 1], S_2d[1, 2], S_2d[2, 2]
                
                # Calculate directional property for 2D
                prop_2d = np.zeros_like(theta_2d)
                prop_name = ""
                prop_unit = "GPa"
                
                for i, t in enumerate(theta_2d):
                    c = np.cos(t)
                    s = np.sin(t)
                    c2, s2 = c**2, s**2
                    c4, s4 = c**4, s**4
                    cs = c * s
                    
                    # Rotated compliance components
                    S_prime_11 = S11*c4 + S22*s4 + (2*S12 + S66)*c2*s2
                    S_prime_22 = S11*s4 + S22*c4 + (2*S12 + S66)*c2*s2
                    S_prime_12 = (S11 + S22 - S66)*c2*s2 + S12*(c4 + s4)
                    S_prime_66 = 4*(S11 + S22 - 2*S12)*c2*s2 + S66*(c2 - s2)**2
                    
                    if 'Young' in selected_prop or selected_prop == 'E':
                        # E(θ) = 1/S'11
                        if abs(S_prime_11) > 1e-12:
                            prop_2d[i] = 1.0 / S_prime_11
                        prop_name = "Young's Modulus E(θ)"
                        prop_unit = "GPa"
                    elif 'Shear' in selected_prop or selected_prop == 'G':
                        # G(θ) = 1/S'66
                        if abs(S_prime_66) > 1e-12:
                            prop_2d[i] = 1.0 / S_prime_66
                        prop_name = "Shear Modulus G(θ)"
                        prop_unit = "GPa"
                    elif 'Poisson' in selected_prop or selected_prop == 'v':
                        # ν(θ) = -S'12/S'11
                        if abs(S_prime_11) > 1e-12:
                            prop_2d[i] = -S_prime_12 / S_prime_11
                        prop_name = "Poisson's Ratio ν(θ)"
                        prop_unit = ""
                    elif 'Bulk' in selected_prop or selected_prop == 'B':
                        # 2D Bulk modulus: B = C11*C22 - C12^2 / (C11 + C22 - 2*C12)
                        # This is a scalar, not directional in 2D
                        B_2d = (C_2d[0,0]*C_2d[1,1] - C_2d[0,1]**2) / (C_2d[0,0] + C_2d[1,1] - 2*C_2d[0,1])
                        prop_2d[i] = B_2d
                        prop_name = "2D Bulk Modulus B"
                        prop_unit = "GPa"
                    else:
                        # Default to Young's modulus
                        if abs(S_prime_11) > 1e-12:
                            prop_2d[i] = 1.0 / S_prime_11
                        prop_name = "Young's Modulus E(θ)"
                        prop_unit = "GPa"
                
                # Create polar plot
                ax = self.canvas.fig.add_subplot(111, projection='polar')
                self.canvas.axes = ax
                ax.plot(theta_2d, prop_2d, 'b-', linewidth=2)
                ax.fill(theta_2d, prop_2d, alpha=0.3)
                title = f"2D {prop_name}"
                if prop_unit:
                    title += f" [{prop_unit}]"
                ax.set_title(title, pad=20)
                ax.set_theta_zero_location('E')  # 0 degrees at right
                ax.set_theta_direction(-1)  # Clockwise
                
                self.canvas.draw()
                if self.window().statusBar():
                    self.window().statusBar().showMessage('2D calculation completed successfully.')
                return
            
            # First, compute compliance matrix (S = C^-1) for 3D materials
            try:
                S = np.linalg.inv(self.cij)
            except np.linalg.LinAlgError:
                QMessageBox.critical(self, 'Error', 'Stiffness matrix is singular! Check your input values.')
                return
            
            # Generate angle meshgrid for 3D plotting
            theta, phi = np.meshgrid(
                np.linspace(0, np.pi, n_points + 1), 
                np.linspace(-np.pi, np.pi, 2 * n_points + 1)
            )
            
            # Direction vectors for plotting
            X = np.sin(theta) * np.cos(phi)
            Y = np.sin(theta) * np.sin(phi)
            Z = np.cos(theta)
            
            # Calculate and plot
            if vis_type == '3D':
                from matplotlib import cm
                from mpl_toolkits.mplot3d import Axes3D
                
                ax = self.canvas.fig.add_subplot(111, projection='3d')
                self.canvas.axes = ax
                
                # Make background transparent to match UI
                ax.set_facecolor('none')
                ax.xaxis.set_pane_color((1.0, 1.0, 1.0, 0.0))
                ax.yaxis.set_pane_color((1.0, 1.0, 1.0, 0.0))
                ax.zaxis.set_pane_color((1.0, 1.0, 1.0, 0.0))
                
                # Apply Grid visibility based on toggle
                show_grid = getattr(self, 'show_grid', True)
                ax.grid(show_grid)
                if not show_grid:
                    ax.axis('off')
                else:
                    ax.axis('on')

                
                # Get selected property from dropdown
                selected_name = self.display_property_combo.currentText()
                prop = self.property_code_map.get(selected_name, 'E')
                
                if prop == 'E':
                    V = Young_3D(S, theta, phi)
                    title = "Young's Modulus (GPa)"
                elif prop == 'G':
                    _, V, _ = Shear_3D(S, theta, phi)
                    title = 'Shear Modulus (GPa)'
                elif prop == 'v':
                    _, V, _ = Poisson_3D(S, theta, phi)
                    title = "Poisson's Ratio"
                elif prop == 'B':
                    V = Bulk_3D(S, theta, phi)
                    title = 'Bulk Modulus (GPa)'
                elif prop == 'H':
                    V = Hardness_3D(S, theta, phi)
                    title = 'Hardness (GPa)'
                else:
                    V = Young_3D(S, theta, phi)
                    title = "Young's Modulus (GPa)"
                
                # Create 3D surface plot
                V_max = np.max(V)
                V_min = np.min(V)
                
                # Create scalar mappable for color mapping
                norm = colors.Normalize(vmin=V_min, vmax=V_max)
                mappable = cm.ScalarMappable(norm=norm, cmap=cm.jet)
                mappable.set_array(V)
                
                # Prepare mesh properties
                show_mesh = getattr(self, 'show_mesh', False)
                lw = 0.3 if show_mesh else 0
                edge_color = 'k' if show_mesh else None 
                
                # Prepare transparency
                show_transparent = getattr(self, 'show_transparent', False)
                surface_alpha = 0.7 if show_transparent else 1.0
                
                surf = ax.plot_surface(
                    V * X, V * Y, V * Z,
                    facecolors=mappable.to_rgba(V),
                    alpha=surface_alpha,
                    linewidth=lw,
                    edgecolor=edge_color,
                    antialiased=True,
                    rcount=n_points,
                    ccount=2*n_points
                )
                
                ax.set_title(f'{title}\nMax: {V_max:.2f}, Min: {V_min:.2f}')
                ax.set_xlabel('X')
                ax.set_ylabel('Y')
                ax.set_zlabel('Z')
                
                # Set default view to Isometric (Applied AFTER plot to ensure it holds)
                ax.view_init(elev=30, azim=-60)
                
                # Set equal aspect ratio (1:1:1)
                # Check for valid V_max before setting limits
                if np.isfinite(V_max) and V_max > 0:
                    max_range = V_max * 1.1
                    ax.set_xlim([-max_range, max_range])
                    ax.set_ylim([-max_range, max_range])
                    ax.set_zlim([-max_range, max_range])
                    ax.set_box_aspect([1, 1, 1])  # Equal aspect ratio
                
                ax.set_box_aspect([1, 1, 1])  # Equal aspect ratio
                
                # Add colorbar using the mappable
                cbar = self.canvas.fig.colorbar(mappable, ax=ax, shrink=0.8, pad=0.1)
                cbar.set_label(f'{title}')
                
            else:
                # 2D plotting with plane selection
                selected_name = self.display_property_combo.currentText()
                prop = self.property_code_map.get(selected_name, 'E')
                
                # Check if using arbitrary plane
                if self.arbitrary_plane_radio.isChecked():
                    h = self.hkl_inputs[0].value()
                    k = self.hkl_inputs[1].value()
                    l = self.hkl_inputs[2].value()
                    slice_plane = [h, k, l]
                    
                    # Verify normal is not zero
                    if np.linalg.norm(slice_plane) < 1e-6:
                        QMessageBox.warning(self, 'Warning', 'Plane normal vector cannot be zero.')
                        return
                    
                    ax = self.canvas.fig.add_subplot(111, projection='polar')
                    self.canvas.axes = ax
                    Plot_Slice_2D(S, n_points, prop, slice_plane, show=False, ax=ax)
                    self.canvas.draw()
                    return
                    
                else:
                    # Standard planes
                    selected_planes = [p for p, cb in self.plane_checkboxes.items() if cb.isChecked()]
                if not selected_planes:
                    selected_planes = ['xy']  # Default to XY plane
                
                n_planes = len(selected_planes)
                
                # Determine subplot layout: use 2x2 for 3 planes, otherwise 1xN
                if n_planes == 3:
                    n_rows, n_cols = 2, 2
                elif n_planes == 2:
                    n_rows, n_cols = 1, 2
                else:
                    n_rows, n_cols = 1, 1
                
                # Create polar subplots for each selected plane
                for idx, plane in enumerate(selected_planes):
                    ax = self.canvas.fig.add_subplot(n_rows, n_cols, idx + 1, projection='polar')
                    self.canvas.axes = ax
                    ax.set_rlabel_position(45)  # Move radial labels to 45 degrees to reduce overlap
                    
                    # Define angles for each plane (1D arrays)
                    t = np.linspace(0, 2 * np.pi, n_points + 1)
                    
                    if plane == 'xy':
                        # XY plane: theta = 90°, phi varies
                        theta_1d = np.full_like(t, np.pi / 2)
                        phi_1d = t
                        plane_label = 'XY Plane (Z=0)'
                    elif plane == 'xz':
                        # XZ plane: phi = 0, theta varies
                        theta_1d = t
                        phi_1d = np.zeros_like(t)
                        plane_label = 'XZ Plane (Y=0)'
                    elif plane == 'yz':
                        # YZ plane: phi = 90°, theta varies
                        theta_1d = t
                        phi_1d = np.full_like(t, np.pi / 2)
                        plane_label = 'YZ Plane (X=0)'
                    else:
                        continue
                    
                    # Calculate property values using direct 1D calculation
                    if prop == 'E':
                        # Young's modulus works with 1D arrays directly
                        V = Young_3D(S, theta_1d, phi_1d)
                        title = "Young's Modulus (GPa)"
                    elif prop == 'G':
                        # Calculate average shear modulus over chi values
                        # Shear_4D is imported at top of file
                        n_chi = 36  # Number of chi angles to average over
                        G_sum = np.zeros_like(t)
                        for chi in np.linspace(-np.pi, np.pi, n_chi):
                            G_sum += Shear_4D(S, theta_1d, phi_1d, chi)
                        V = G_sum / n_chi  # Average shear modulus
                        title = 'Shear Modulus (GPa)'
                    elif prop == 'v':
                        # Calculate average Poisson's ratio over chi values
                        # Poisson_4D is imported at top of file
                        n_chi = 36
                        v_sum = np.zeros_like(t)
                        for chi in np.linspace(-np.pi, np.pi, n_chi):
                            v_sum += Poisson_4D(S, theta_1d, phi_1d, chi)
                        V = v_sum / n_chi  # Average Poisson's ratio
                        title = "Poisson's Ratio"
                    elif prop == 'B':
                        # Bulk modulus works with 1D arrays directly
                        V = Bulk_3D(S, theta_1d, phi_1d)
                        title = 'Bulk Modulus (GPa)'
                    elif prop == 'H':
                        # Calculate Hardness using average shear modulus
                        # Shear_4D is imported at top of file
                        n_chi = 36
                        G_sum = np.zeros_like(t)
                        for chi in np.linspace(-np.pi, np.pi, n_chi):
                            G_sum += Shear_4D(S, theta_1d, phi_1d, chi)
                        G_avg = G_sum / n_chi
                        E = Young_3D(S, theta_1d, phi_1d)
                        k = G_avg / E  # Pugh's ratio
                        k2G = k * k * G_avg
                        k2G_safe = np.maximum(k2G, 1e-10)
                        V = 2 * (k2G_safe ** 0.585) - 3
                        V = np.nan_to_num(V, nan=0.0, posinf=0.0, neginf=0.0)
                        title = 'Hardness (GPa)'
                    else:
                        V = Young_3D(S, theta_1d, phi_1d)
                        title = "Young's Modulus (GPa)"
                    
                    # Handle NaN values
                    V = np.nan_to_num(V, nan=0.0, posinf=0.0, neginf=0.0)
                    
                    # Plot
                    ax.plot(t, V, 'r-', linewidth=2)
                    ax.fill(t, V, alpha=0.3, color='red')
                    ax.set_title(f'{title}\n{plane_label}')
                
                # Adjust layout for better spacing
                self.canvas.fig.tight_layout()
            
            # Update the canvas
            self.canvas.draw()
            
            # Calculate and display VRH averages
            vrh_results = ElasticVRH3D(self.cij)
            status_message = f"VRH Averages: E={vrh_results['E']:.2f} GPa, G={vrh_results['G_VRH']:.2f} GPa, v={vrh_results['v']:.4f}, B={vrh_results['K_VRH']:.2f} GPa"
            if self.window().statusBar():
                self.window().statusBar().showMessage(status_message)
            
            # Check mechanical stability
            stability_res = check_stability_detailed(self.cij)
            if stability_res['stable']:
                self.stability_label.setText("The Structure is STABLE")
                self.stability_label.setStyleSheet("padding-right: 20px; font-weight: bold; color: green;")
            else:
                self.stability_label.setText("The Structure is UNSTABLE")
                self.stability_label.setStyleSheet("padding-right: 20px; font-weight: bold; color: red;")
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, 'Error', f'Calculation error: {str(e)}')
    
    def _fillCijBySymmetry(self):
        """Fill dependent elastic constants based on crystal type symmetry."""
        crystal_type = self.crystal_type_combo.currentText()
        
        if crystal_type == 'Cubic':
            # C11 = C22 = C33
            self.cij[1, 1] = self.cij[0, 0]  # C22 = C11
            self.cij[2, 2] = self.cij[0, 0]  # C33 = C11
            # C12 = C13 = C23
            self.cij[0, 2] = self.cij[0, 1]  # C13 = C12
            self.cij[2, 0] = self.cij[0, 1]  # C31 = C12
            self.cij[1, 2] = self.cij[0, 1]  # C23 = C12
            self.cij[2, 1] = self.cij[0, 1]  # C32 = C12
            # C44 = C55 = C66
            self.cij[4, 4] = self.cij[3, 3]  # C55 = C44
            self.cij[5, 5] = self.cij[3, 3]  # C66 = C44
            
        elif crystal_type == 'Hexagonal':
            # C11 = C22
            self.cij[1, 1] = self.cij[0, 0]  # C22 = C11
            # C13 = C23
            self.cij[1, 2] = self.cij[0, 2]  # C23 = C13
            self.cij[2, 1] = self.cij[0, 2]  # C32 = C13
            # C44 = C55
            self.cij[4, 4] = self.cij[3, 3]  # C55 = C44
            # C66 = (C11 - C12) / 2
            self.cij[5, 5] = (self.cij[0, 0] - self.cij[0, 1]) / 2
            
        elif crystal_type == 'Tetragonal':
            # C11 = C22
            self.cij[1, 1] = self.cij[0, 0]  # C22 = C11
            # C13 = C23
            self.cij[1, 2] = self.cij[0, 2]  # C23 = C13
            self.cij[2, 1] = self.cij[0, 2]  # C32 = C13
            # C44 = C55
            self.cij[4, 4] = self.cij[3, 3]  # C55 = C44
            
        elif crystal_type == 'Orthorhombic':
            # No additional symmetry constraints, all 9 constants independent
            pass
            
        elif crystal_type == 'Monoclinic':
            # No additional symmetry constraints (13 independent constants)
            pass
            
        elif crystal_type == 'Triclinic':
            # No symmetry constraints (21 independent constants)
            pass
        
        # Ensure matrix is symmetric
        for i in range(6):
            for j in range(i+1, 6):
                self.cij[j, i] = self.cij[i, j]

# Removed createMenuBar

    def exportData(self):
        """Export elastic model to external software formats"""
        # Check if we have data
        if np.allclose(self.cij, 0):
             QMessageBox.warning(self, 'Warning', 'No data to export. Please input/load elastic constants first.')
             return
             
        # Create dialog for export
        # Option: "ABAQUS UMAT (*.f)", "COMSOL Parameters (*.txt)"
        file_path, filter = QFileDialog.getSaveFileName(
            self, 'Export Material Model', '',
            'ABAQUS UMAT (*.f);;COMSOL Parameters (*.txt)'
        )
        
        if not file_path:
            return
            
        try:
            # Determine format
            if 'ABAQUS' in filter:
                fmt = 'abaqus'
                if not file_path.lower().endswith('.f'):
                    file_path += '.f'
            elif 'COMSOL' in filter:
                fmt = 'comsol'
                if not file_path.lower().endswith('.txt'):
                    file_path += '.txt'
            else:
                return

            # Get material info
            mat_name = self.data_name if hasattr(self, 'data_name') else "Custom Material"
            crystal_type = self.crystal_type_combo.currentText()
            
            # Export
            export_model(file_path, fmt, self.cij, mat_name, crystal_type)
            
            QMessageBox.information(self, 'Success', f'Model exported successfully to:\n{file_path}')
            
        except Exception as e:
            QMessageBox.critical(self, 'Error', f'Export failed: {str(e)}')

    def generateReport(self, report_format='html'):
        """Generate and save an analysis report (html, pdf, or docx)"""
        # data check
        if np.allclose(self.cij, 0):
             QMessageBox.warning(self, 'Warning', 'No data to report. Please input/load elastic constants first.')
             return
             
        # Ask user where to save - unified dialog with all format options
        file_filter = "HTML Files (*.html);;PDF Files (*.pdf);;Word Documents (*.docx)"
        file_path, selected_filter = QFileDialog.getSaveFileName(self, "Save Report", "", file_filter)
        
        if not file_path:
            return
        
        # Determine format from selected filter
        if "*.pdf" in selected_filter:
            report_format = 'pdf'
            ext = '.pdf'
        elif "*.docx" in selected_filter:
            report_format = 'docx'
            ext = '.docx'
        else:
            report_format = 'html'
            ext = '.html'
            
        if not file_path.lower().endswith(ext):
            file_path += ext
            
        # Store current property to restore later
        current_property = self.display_property_combo.currentText()
        temp_files = []
        plots = {}
        
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            import tempfile

            
            # List of properties to plot
            properties = [
                "Young's Modulus (E)",
                "Shear Modulus (G)",
                "Poisson's Ratio (v)",
                "Bulk Modulus (B)", # Added these to match previous code
                "Hardness (H)"
            ]
            
            # Check availability in combo before selecting
            available_props = [self.display_property_combo.itemText(i) for i in range(self.display_property_combo.count())]
            
            for prop_name in properties:
                if prop_name not in available_props:
                    continue
                    
                # Switch property
                self.display_property_combo.blockSignals(True)
                self.display_property_combo.setCurrentText(prop_name)
                self.display_property_combo.blockSignals(False)
                
                # Force recalculate/replot
                self._doCalculation()
                QApplication.processEvents() # Ensure update
                
                # Save plot
                fd, temp_img_path = tempfile.mkstemp(suffix='.png')
                os.close(fd)
                temp_files.append(temp_img_path)
                
                # Use figure save to get high quality
                self.canvas.fig.savefig(temp_img_path, dpi=100, bbox_inches='tight')
                plots[prop_name] = temp_img_path
            
            # Create report
            # Get material info
            mat_name = self.data_name if hasattr(self, 'data_name') else "Custom Material"
            if hasattr(self, 'data_source') and self.data_source:
                source_name = os.path.basename(self.data_source).split('.')[0]
                if source_name != mat_name:
                    mat_name = source_name
            
            # Determine target HTML path
            if report_format == 'pdf':
                fd, target_html = tempfile.mkstemp(suffix='.html')
                os.close(fd)
                temp_files.append(target_html)
            else:
                target_html = file_path
                
            # Generate HTML
            final_html_path = generate_report(
                filename=target_html,
                cij_matrix=self.cij,
                material_name=mat_name,
                crystal_type=self.crystal_type_combo.currentText(),
                plots=plots
            )
            
            if report_format == 'pdf':
                # Convert to PDF
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(file_path)
                # printer.setPageSize(QPrinter.PageSize.A4) # Enum might vary by Qt version, default is usually A4
                
                doc = QTextDocument()
                with open(final_html_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                    doc.setHtml(html_content)
                doc.print(printer)
                
            elif report_format == 'docx':
                # Convert to Word using python-docx
                try:
                    from docx import Document
                    from docx.shared import Inches, Pt, Cm
                    from docx.enum.section import WD_ORIENT
                except ImportError:
                    QMessageBox.warning(self, 'Missing Dependency', 
                        'Word export requires python-docx package.\n'
                        'Please install it with: pip install python-docx')
                    return
                
                # Create Word document
                doc = Document()
                
                # Set A4 page size
                section = doc.sections[0]
                section.page_width = Cm(21.0)  # A4 width
                section.page_height = Cm(29.7)  # A4 height
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                
                # Add title
                title = doc.add_heading(f'Mat Model Lab Analysis Report', level=0)
                
                # Add material info
                doc.add_heading('Material Information', level=1)
                doc.add_paragraph(f'Material Name: {mat_name}')
                doc.add_paragraph(f'Crystal Type: {self.crystal_type_combo.currentText()}')
                
                # Add Cij matrix
                doc.add_heading('Elastic Stiffness Matrix (GPa)', level=1)
                table = doc.add_table(rows=7, cols=7)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = ''
                for j in range(6):
                    header_cells[j+1].text = f'C{j+1}'
                
                # Data rows
                for i in range(6):
                    row_cells = table.rows[i+1].cells
                    row_cells[0].text = f'C{i+1}'
                    for j in range(6):
                        val = self.cij[i, j]
                        row_cells[j+1].text = f'{val:.2f}' if val != 0 else '-'
                
                # Add plots
                doc.add_heading('Visualizations', level=1)
                for prop_name, img_path in plots.items():
                    doc.add_paragraph(prop_name)
                    # Add image with width of 15cm (fits A4 with margins)
                    doc.add_picture(img_path, width=Cm(15))
                
                # Save document
                doc.save(file_path)
                
            QApplication.restoreOverrideCursor()
            QMessageBox.information(self, 'Success', f'Report generated successfully:\n{file_path}')
            
            # Open if requested (only for HTML usually, but PDF works too if system association exists)
            # import webbrowser
            # webbrowser.open('file://' + file_path.replace('\\', '/'))

        except Exception as e:
            import traceback
            traceback.print_exc()
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(self, 'Error', f'Failed to generate report:\n{str(e)}')
            
        finally:
            # Restore state
            self.display_property_combo.blockSignals(True)
            self.display_property_combo.setCurrentText(current_property)
            self.display_property_combo.blockSignals(False)
            self._doCalculation()
            
            # Clean up temp files
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                    except:
                        pass
            
            QApplication.restoreOverrideCursor()


    def showDocumentation(self):
        """Show the documentation dialog"""
        from .widgets.help_dialog import HelpDialog
        # Check if current_theme is dark
        is_dark = getattr(self, 'current_theme', 'light') == 'dark'
        dialog = HelpDialog(self, is_dark=is_dark)
        dialog.exec()
        
    def showAbout(self):
        """Show about dialog"""
        QMessageBox.about(self, "About Mat Model Lab",
            "Mat Model Lab\n\n"
            "A comprehensive material constitutive model analysis and visualization tool.\n\n"
            "Version 1.0\n"
            "(c) 2026")

    def toggleGrid(self, checked):
        """Toggle grid visibility for 3D plots"""
        self.show_grid = checked
        if hasattr(self, 'canvas'):
            if self.vis_3d_radio.isChecked():
                self._doCalculation()
            else:
                 self._doCalculation()

    def toggleMesh(self, checked):
        """Toggle surface mesh visibility for 3D plots"""
        self.show_mesh = checked
        if hasattr(self, 'canvas'):
             if self.vis_3d_radio.isChecked():
                self._doCalculation()

    def toggleTransparency(self, checked):
        """Toggle surface transparency for 3D plots"""
        self.show_transparent = checked
        if hasattr(self, 'canvas'):
            if self.vis_3d_radio.isChecked():
                self._doCalculation()
                
    def _create_toolbar_icon(self, name, color):
        """Create a simple icon for toolbar buttons"""
        pixmap = QPixmap(32, 32)
        pixmap.fill(Qt.GlobalColor.transparent)
        
        painter = QPainter(pixmap)
        pen = QPen(QColor(color))
        pen.setWidth(2)
        painter.setPen(pen)
        
        if name == 'grid':
            # Draw a grid
            painter.drawRect(4, 4, 24, 24)
            painter.drawLine(12, 4, 12, 28)
            painter.drawLine(20, 4, 20, 28)
            painter.drawLine(4, 12, 28, 12)
            painter.drawLine(4, 20, 28, 20)
            
        elif name == 'mesh':
            # Draw a wireframe sphere-like shape
            painter.drawEllipse(4, 4, 24, 24)
            # Longitude
            painter.drawArc(4, 4, 24, 24, 0, 16 * 360) # Outer circle
            painter.drawEllipse(10, 4, 12, 24) # Inner ellipse
            # Latitude
            painter.drawLine(4, 16, 28, 16) # Equator
            # painter.drawArc(4, 10, 24, 12, 0, 16 * 180) # Upper arc, tricky with drawArc
        
        elif name == 'alpha':
            # Draw a half-filled/half-empty square to represent transparency
            # Left half filled (semi-transparent feel)
            from PyQt6.QtCore import QRect
            painter.drawRect(4, 4, 24, 24)
            # Draw a gradient from opaque (left) to empty (right) using lines
            # Or simpler: draw a checkboard pattern (transparency pattern)
            # Filled left half
            brush = painter.brush()
            from PyQt6.QtGui import QBrush
            painter.setBrush(QBrush(QColor(color)))
            painter.drawRect(4, 4, 12, 24)  # Left filled
            painter.setBrush(brush)  # Reset
            # Diagonal line as a separator
            painter.drawLine(16, 4, 16, 28)
        
        painter.end()
        return QIcon(pixmap)

